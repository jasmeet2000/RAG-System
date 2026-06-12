"""
Document Loaders — the entry point for all documents into the RAG system.

=== WHY THIS FILE EXISTS ===
Before we can search or answer questions about a document, we need to read it.
Different file formats require completely different libraries and parsing logic.
This module abstracts all format-specific loading behind a single, unified interface.

=== HOW IT INTERACTS WITH OTHER MODULES ===
  Caller:  app/ingestion/parser.py  (calls load_document())
  Output:  RawDocument dataclass → passed to parser → cleaner → chunker → embeddings
  Uses:    app/core/exceptions.py  (raises DocumentLoadError)
           app/core/constants.py   (SUPPORTED_FILE_EXTENSIONS, MIME_TYPE_MAP)
           app/core/logging.py     (structured logging)

=== WHAT HAPPENS WITHOUT IT ===
The pipeline has no way to read files. Every other module assumes text has
already been extracted. Without the loader, the entire ingestion pipeline is broken.

=== INDUSTRY ALTERNATIVES ===
- LangChain's DocumentLoader: Abstracts the same thing, but hides the implementation.
  We build it ourselves so you understand what PyMuPDF, python-docx etc. actually do.
- Apache Tika: Java-based universal parser, overkill for our use case.
- Unstructured.io: Commercial-grade parser (handles tables, images, forms) — excellent
  but complex. Our loaders handle the 80% use case cleanly.
- Docling (IBM): Newer, handles complex PDF layouts well. Good alternative to PyMuPDF.

=== DESIGN PATTERN USED ===
Strategy Pattern via an abstract base class (DocumentLoader).
Each file type gets its own concrete loader (PDFLoader, DOCXLoader, etc.).
The factory function load_document() selects the right loader based on file extension.

Why Strategy Pattern?
- Open/Closed Principle: Add new formats (HTML, CSV) without modifying existing loaders
- Testability: Each loader can be tested independently
- Single Responsibility: Each loader knows about exactly one file format
"""

from __future__ import annotations

import time
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional, Union

from app.core.constants import SUPPORTED_FILE_EXTENSIONS
from app.core.exceptions import DocumentLoadError
from app.core.logging import get_logger

logger = get_logger(__name__)


# ── Data Models ────────────────────────────────────────────────────────────────


@dataclass
class RawDocument:
    """
    Represents a document after loading — before parsing or chunking.

    Why a dataclass instead of a dict?
    - Type safety: IDEs and mypy catch mistakes at development time
    - Documentation: Field names and types are self-documenting
    - Immutability option: Can use frozen=True for read-only documents

    Fields:
        content:      The full extracted text of the document.
        source:       Absolute path to the original file.
        filename:     Just the filename (e.g., "report.pdf").
        file_type:    Extension (e.g., ".pdf").
        page_count:   Number of pages (None for non-paginated formats like TXT).
        file_size_bytes: Size of the original file.
        loaded_at:    UTC timestamp of when the document was loaded.
        metadata:     Any additional loader-specific metadata (page numbers, etc.).
    """

    content: str
    source: str
    filename: str
    file_type: str
    page_count: Optional[int] = None
    file_size_bytes: int = 0
    loaded_at: str = field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    metadata: dict = field(default_factory=dict)

    @property
    def char_count(self) -> int:
        """Total character count of the document content."""
        return len(self.content)

    @property
    def word_count(self) -> int:
        """Approximate word count (split on whitespace)."""
        return len(self.content.split())


# ── Abstract Base Class ────────────────────────────────────────────────────────


class DocumentLoader(ABC):
    """
    Abstract base class for all document loaders.

    Why ABC?
    - Enforces that every concrete loader implements the load() method
    - Provides a common interface: all loaders → RawDocument
    - Enables polymorphism: load_document() can work with any loader

    Industry note: This is the same pattern LangChain uses internally,
    just without the abstraction framework on top.
    """

    @abstractmethod
    def load(self, file_path: Path) -> RawDocument:
        """
        Load a document from disk and return its text content + metadata.

        Args:
            file_path: Absolute path to the document.

        Returns:
            RawDocument with extracted text and metadata.

        Raises:
            DocumentLoadError: If the file cannot be read or parsed.
        """
        ...

    def _validate_file(self, file_path: Path) -> None:
        """
        Common validation shared by all loaders.

        Checks existence and readability before attempting to parse.
        This prevents confusing errors deep in the parsing libraries.
        """
        if not file_path.exists():
            raise DocumentLoadError(
                f"File not found: {file_path}",
                detail="Ensure the file path is correct and the file exists.",
            )
        if not file_path.is_file():
            raise DocumentLoadError(
                f"Path is not a file: {file_path}",
                detail="Directories cannot be loaded. Provide a file path.",
            )
        if file_path.stat().st_size == 0:
            raise DocumentLoadError(
                f"File is empty: {file_path.name}",
                detail="Cannot process empty files.",
            )


# ── Concrete Loaders ───────────────────────────────────────────────────────────


class PDFLoader(DocumentLoader):
    """
    PDF document loader using PyMuPDF (fitz).

    WHY PyMuPDF over alternatives?
    - PyPDF2: Slow, fragile with real-world PDFs, deprecated in favor of pypdf
    - pypdf: Better than PyPDF2, but still struggles with complex layouts
    - pdfminer.six: Good for complex layouts, very slow, difficult API
    - PyMuPDF: 10-50x faster than alternatives, handles complex layouts,
                extracts text in reading order, preserves metadata (author, title, dates)

    PRODUCTION CONSIDERATION:
    PDFs are deceptively complex. A "simple" PDF might be:
    - A scanned image (no text layer) → needs OCR (Tesseract, AWS Textract)
    - Multi-column academic paper → reading order is tricky
    - Password-protected → needs password handling
    - Form with embedded fields → PyMuPDF handles these well

    We handle the common case. For production, consider Unstructured.io or Docling
    for PDFs with tables, images, and complex layouts.
    """

    def load(self, file_path: Path) -> RawDocument:
        """
        Extract text from all pages of a PDF.

        Text from each page is joined with a double newline, preserving the
        page boundary as a logical document structure cue.
        """
        self._validate_file(file_path)

        try:
            import fitz  # PyMuPDF — imported here to avoid import error if not installed
        except ImportError as e:
            raise DocumentLoadError(
                "PyMuPDF is not installed. Run: pip install PyMuPDF",
                detail=str(e),
            ) from e

        start_time = time.perf_counter()

        try:
            doc = fitz.open(str(file_path))
            page_count = len(doc)

            pages_text: list[str] = []
            for page_num, page in enumerate(doc, start=1):
                text = page.get_text("text")  # "text" mode = plain text in reading order
                if text.strip():  # Skip empty pages
                    pages_text.append(text)
                else:
                    logger.debug(
                        "Skipping empty PDF page",
                        file=file_path.name,
                        page=page_num,
                    )

            # Extract PDF metadata (author, title, creation date, etc.)
            pdf_metadata = doc.metadata or {}
            doc.close()

        except Exception as e:
            raise DocumentLoadError(
                f"Failed to parse PDF: {file_path.name}",
                detail=str(e),
            ) from e

        full_text = "\n\n".join(pages_text)
        elapsed = time.perf_counter() - start_time

        logger.info(
            "PDF loaded",
            file=file_path.name,
            pages=page_count,
            chars=len(full_text),
            elapsed_ms=round(elapsed * 1000, 2),
        )

        return RawDocument(
            content=full_text,
            source=str(file_path.resolve()),
            filename=file_path.name,
            file_type=".pdf",
            page_count=page_count,
            file_size_bytes=file_path.stat().st_size,
            metadata={
                "pdf_title": pdf_metadata.get("title", ""),
                "pdf_author": pdf_metadata.get("author", ""),
                "pdf_created": pdf_metadata.get("creationDate", ""),
                "pdf_producer": pdf_metadata.get("producer", ""),
            },
        )


class DOCXLoader(DocumentLoader):
    """
    DOCX document loader using python-docx.

    DOCX structure:
    A .docx file is a ZIP archive containing XML files. python-docx parses
    this XML and exposes Paragraphs, Tables, Styles, etc. through a clean API.

    What we extract:
    - Paragraphs (the main body text)
    - Table cells (often contain important structured data)
    - Heading levels (preserved as text for context)

    What we intentionally skip:
    - Images (would need vision models to process)
    - Headers/footers (usually boilerplate — page numbers, company names)
    - Footnotes (context-specific — enable if your docs use them heavily)

    PRODUCTION CONSIDERATION:
    DOCX documents often have rich structure (headings, tables, bullet points).
    A more sophisticated loader could use heading levels to guide chunking,
    e.g., keeping H1 → H2 → paragraph as a semantic unit.
    """

    def load(self, file_path: Path) -> RawDocument:
        """Extract text from paragraphs and tables in a DOCX file."""
        self._validate_file(file_path)

        try:
            from docx import Document  # python-docx
        except ImportError as e:
            raise DocumentLoadError(
                "python-docx is not installed. Run: pip install python-docx",
                detail=str(e),
            ) from e

        start_time = time.perf_counter()

        try:
            doc = Document(str(file_path))
            text_blocks: list[str] = []

            # Extract paragraphs — the main text body
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    text_blocks.append(text)

            # Extract table cell content
            # Tables are separate from the paragraph flow in DOCX
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_blocks.append(row_text)

            # Extract document-level metadata
            core_props = doc.core_properties
            doc_metadata = {
                "docx_title": core_props.title or "",
                "docx_author": core_props.author or "",
                "docx_created": str(core_props.created or ""),
                "docx_modified": str(core_props.modified or ""),
                "docx_subject": core_props.subject or "",
            }

        except Exception as e:
            raise DocumentLoadError(
                f"Failed to parse DOCX: {file_path.name}",
                detail=str(e),
            ) from e

        full_text = "\n\n".join(text_blocks)
        elapsed = time.perf_counter() - start_time

        if not full_text.strip():
            raise DocumentLoadError(
                f"No text extracted from DOCX: {file_path.name}",
                detail="The document may be empty, image-only, or corrupted.",
            )

        logger.info(
            "DOCX loaded",
            file=file_path.name,
            paragraphs=len(text_blocks),
            chars=len(full_text),
            elapsed_ms=round(elapsed * 1000, 2),
        )

        return RawDocument(
            content=full_text,
            source=str(file_path.resolve()),
            filename=file_path.name,
            file_type=".docx",
            page_count=None,  # DOCX doesn't have a fixed page count concept
            file_size_bytes=file_path.stat().st_size,
            metadata=doc_metadata,
        )


class TXTLoader(DocumentLoader):
    """
    Plain text document loader.

    The simplest loader — reads the file as UTF-8 text.

    Encoding strategy:
    We try UTF-8 first (modern standard), then fall back to latin-1 (legacy).
    latin-1 is a "safe" fallback because it can decode any byte sequence
    (every byte 0-255 maps to a valid character), even if some characters
    look garbled. This is better than crashing on encoding errors.

    PRODUCTION CONSIDERATION:
    In real-world systems, you'll encounter files with mixed encodings.
    The `chardet` library can auto-detect encoding — add it if you process
    documents from diverse sources (scraped web content, legacy systems).
    """

    ENCODING_FALLBACKS: list[str] = ["utf-8", "utf-8-sig", "latin-1"]

    def load(self, file_path: Path) -> RawDocument:
        """Read plain text file with encoding auto-detection fallback."""
        self._validate_file(file_path)

        start_time = time.perf_counter()
        content: str = ""
        used_encoding: str = ""

        for encoding in self.ENCODING_FALLBACKS:
            try:
                content = file_path.read_text(encoding=encoding)
                used_encoding = encoding
                break
            except UnicodeDecodeError:
                logger.debug(
                    "Encoding failed, trying next",
                    file=file_path.name,
                    tried_encoding=encoding,
                )
                continue

        if not content:
            raise DocumentLoadError(
                f"Could not decode text file: {file_path.name}",
                detail=f"Tried encodings: {self.ENCODING_FALLBACKS}",
            )

        elapsed = time.perf_counter() - start_time

        logger.info(
            "TXT loaded",
            file=file_path.name,
            encoding=used_encoding,
            chars=len(content),
            elapsed_ms=round(elapsed * 1000, 2),
        )

        return RawDocument(
            content=content,
            source=str(file_path.resolve()),
            filename=file_path.name,
            file_type=".txt",
            page_count=None,
            file_size_bytes=file_path.stat().st_size,
            metadata={"encoding": used_encoding},
        )


class MarkdownLoader(DocumentLoader):
    """
    Markdown document loader.

    Two-step process:
    1. Read raw markdown text (for chunking purposes, raw markdown is often fine)
    2. Optionally strip markdown syntax for cleaner embedding text

    Why keep markdown syntax by default?
    - Headers (#, ##) signal document structure — useful for parent-child chunking
    - Code blocks (```) are meaningful context
    - **Bold** and *italic* text emphasizes important terms

    The cleaner module (cleaner.py) can strip markdown syntax if needed.
    We load it raw here and let downstream stages decide.
    """

    def load(self, file_path: Path) -> RawDocument:
        """Read markdown file preserving structure markers."""
        self._validate_file(file_path)

        start_time = time.perf_counter()

        # Markdown is just text — use TXTLoader's robust encoding handling
        txt_loader = TXTLoader()
        raw_doc = txt_loader.load(file_path)

        elapsed = time.perf_counter() - start_time

        logger.info(
            "Markdown loaded",
            file=file_path.name,
            chars=raw_doc.char_count,
            elapsed_ms=round(elapsed * 1000, 2),
        )

        # Override the file_type to distinguish .md from .txt
        return RawDocument(
            content=raw_doc.content,
            source=raw_doc.source,
            filename=raw_doc.filename,
            file_type=".md",
            page_count=None,
            file_size_bytes=raw_doc.file_size_bytes,
            metadata={**raw_doc.metadata, "format": "markdown"},
        )


# ── Factory Function ──────────────────────────────────────────────────────────


# Maps file extensions to their loader classes.
# Adding a new format: add entry here + create a new Loader class above.
_LOADER_REGISTRY: dict[str, type[DocumentLoader]] = {
    ".pdf": PDFLoader,
    ".docx": DOCXLoader,
    ".txt": TXTLoader,
    ".md": MarkdownLoader,
}


def load_document(file_path: Union[str, Path]) -> RawDocument:
    """
    Load a document from disk by selecting the appropriate loader automatically.

    This is the public API of the loader module. All callers use this function
    rather than instantiating loaders directly.

    WHY A FACTORY FUNCTION?
    - Callers don't need to know which loader to use — they just pass a file path
    - Adding a new format only requires updating _LOADER_REGISTRY
    - Keeps the interface stable even as we add new loaders

    This implements the Factory Method pattern: a function that creates the
    right object based on runtime input.

    Args:
        file_path: Path to the document (str or Path).

    Returns:
        RawDocument containing the extracted text and metadata.

    Raises:
        DocumentLoadError: If the file format is unsupported or loading fails.

    Example:
        doc = load_document("data/raw/annual_report.pdf")
        print(doc.word_count)  # 15234
    """
    path = Path(file_path)
    extension = path.suffix.lower()

    if extension not in SUPPORTED_FILE_EXTENSIONS:
        raise DocumentLoadError(
            f"Unsupported file format: '{extension}' ({path.name})",
            detail=(
                f"Supported formats: {', '.join(sorted(SUPPORTED_FILE_EXTENSIONS))}. "
                "To add support for a new format, create a new DocumentLoader subclass."
            ),
        )

    loader_class = _LOADER_REGISTRY[extension]
    loader = loader_class()

    logger.debug(
        "Loading document",
        file=path.name,
        format=extension,
        loader=loader_class.__name__,
    )

    return loader.load(path)


def get_supported_extensions() -> set[str]:
    """Return the set of file extensions this loader module supports."""
    return set(_LOADER_REGISTRY.keys())
